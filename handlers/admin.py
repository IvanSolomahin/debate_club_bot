# handlers/admin.py
from datetime import datetime, timedelta
from aiogram import Router, F, types
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import CallbackQuery, Message, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.utils.keyboard import InlineKeyboardBuilder
from sqlalchemy.ext.asyncio import AsyncSession
from aiogram import Bot

import repo
from db import Training, User
from config import settings


router = Router()


# ── States ──────────────────────────────────────────────────────────────────

class AdminEventStates(StatesGroup):
    waiting_title = State()
    waiting_datetime = State()
    waiting_place = State()
    waiting_description = State()
    waiting_confirm = State()
    waiting_new_value = State()


class AdminBroadcastStates(StatesGroup):
    waiting_text = State()


class AdminManageAdminStates(StatesGroup):
    waiting_user_id = State()


# ── Date/Time Helpers ───────────────────────────────────────────────────────

DATE_FORMAT_DISPLAY = "%d.%m %H:%M"


def _generate_date_keyboard() -> InlineKeyboardMarkup:
    """Generate keyboard with nearest dates (next 7 days, every day at 10:00, 14:00, 18:00)"""
    builder = InlineKeyboardBuilder()
    now = datetime.now()
    
    # Generate options for the next 7 days with common time slots
    time_slots = [(18,30)]  # 10:00, 14:00, 18:00
    
    for day_offset in range(7):
        date = now + timedelta(days=day_offset)
        for hour, minute in time_slots:
            dt = date.replace(hour=hour, minute=minute, second=0, microsecond=0)
            display = dt.strftime(DATE_FORMAT_DISPLAY)
            callback_data = f"date_select:{dt.strftime('%Y%m%d%H%M')}"
            builder.button(text=display, callback_data=callback_data)
    
    builder.button(text="✏️ Ввести вручную", callback_data="date_manual")
    builder.adjust(2)  # 2 buttons per row
    return builder.as_markup()


def _parse_datetime_manual(text: str) -> datetime | None:
    """Parse datetime in DD.MM HH:MM format"""
    try:
        # Try DD.MM HH:MM format (assumes current year)
        dt = datetime.strptime(text, "%d.%m %H:%M")
        # Set year to current since strptime defaults to 1900
        dt = dt.replace(year=datetime.now().year)
        return dt
    except ValueError:
        return None


# ── UC-06 Create event ──────────────────────────────────────────────────────

@router.callback_query(F.data == "admin_create_event")
async def cmd_create_event(callback: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(AdminEventStates.waiting_title)
    await callback.message.answer("Введите название события:")
    await callback.answer()


@router.message(AdminEventStates.waiting_title)
async def input_title(message: Message, state: FSMContext) -> None:
    if message.text is None:
        return
    await state.update_data(title=message.text)
    await state.set_state(AdminEventStates.waiting_datetime)
    await message.answer(
        "Выберите дату и время события:",
        reply_markup=_generate_date_keyboard()
    )


@router.callback_query(F.data.startswith("date_select:"), AdminEventStates.waiting_datetime)
async def select_datetime_from_keyboard(callback: CallbackQuery, state: FSMContext) -> None:
    """Handle date selection from keyboard"""
    date_str = callback.data.split(":")[1]  # Format: YYYYMMDDHHMM
    try:
        dt = datetime.strptime(date_str, "%Y%m%d%H%M")
    except ValueError:
        await callback.answer("Неверный формат даты", show_alert=True)
        return

    await state.update_data(datetime=dt.isoformat())
    await state.set_state(AdminEventStates.waiting_place)
    await callback.message.answer(f"Выбрано: {dt.strftime(DATE_FORMAT_DISPLAY)}\nВведите место проведения:")
    await callback.answer()


@router.callback_query(F.data == "date_manual", AdminEventStates.waiting_datetime)
async def request_manual_datetime(callback: CallbackQuery, state: FSMContext) -> None:
    """Request manual datetime input"""
    await callback.message.answer("Введите дату и время в формате ДД.ММ ЧЧ:ММ (например: 15.01 14:00):")
    await callback.answer()


@router.message(AdminEventStates.waiting_datetime)
async def input_datetime(message: Message, state: FSMContext) -> None:
    if message.text is None:
        return
    
    dt = _parse_datetime_manual(message.text)
    if dt is None:
        await message.answer("Неверный формат даты. Используйте формат ДД.ММ ЧЧ:ММ (например: 15.01 14:00)")
        return

    await state.update_data(datetime=dt.isoformat())
    await state.set_state(AdminEventStates.waiting_place)
    await message.answer("Введите место проведения:")


@router.message(AdminEventStates.waiting_place)
async def input_place(message: Message, state: FSMContext) -> None:
    if message.text is None:
        return
    await state.update_data(place=message.text)
    await state.set_state(AdminEventStates.waiting_description)
    await message.answer("Введите описание события:")


@router.message(AdminEventStates.waiting_description)
async def input_description(message: Message, state: FSMContext, session: AsyncSession) -> None:
    if message.text is None or message.from_user is None:
        return
    
    data = await state.get_data()
    description = message.text

    # Convert datetime string back to datetime object
    dt = datetime.fromisoformat(data["datetime"])

    await repo.create_training(
        session=session,
        title=data["title"],
        dt=dt,
        created_by=message.from_user.id,
        description=description,
        location=data["place"],
    )

    await state.clear()
    await message.answer("Событие создано ✅")


# ── UC-07 Edit event ────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin_edit_event")
async def cmd_edit_event(callback: CallbackQuery, session: AsyncSession) -> None:
    trainings = await repo.get_upcoming_trainings(session)

    if not trainings:
        await callback.message.answer("Нет предстоящих событий для редактирования")
        await callback.answer()
        return

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t.title, callback_data=f"edit_select:{t.id}")]
        for t in trainings
    ])

    await callback.message.answer("Выберите событие для редактирования:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("edit_select:"))
async def select_event_to_edit(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.data is None:
        return
    event_id = int(callback.data.split(":")[1])
    await state.update_data(event_id=event_id)

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Название", callback_data="edit_field:title")],
        [InlineKeyboardButton(text="Дата/время", callback_data="edit_field:datetime")],
        [InlineKeyboardButton(text="Место", callback_data="edit_field:place")],
        [InlineKeyboardButton(text="Описание", callback_data="edit_field:description")],
    ])

    await callback.message.answer("Выберите поле для редактирования:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("edit_field:"))
async def select_field_to_edit(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.data is None:
        return
    field = callback.data.split(":")[1]
    await state.update_data(field=field)
    await state.set_state(AdminEventStates.waiting_new_value)
    
    if field == "datetime":
        await callback.message.answer(
            "Выберите новую дату и время:",
            reply_markup=_generate_date_keyboard()
        )
    else:
        await callback.message.answer("Введите новое значение:")
    await callback.answer()


@router.callback_query(F.data.startswith("date_select:"), AdminEventStates.waiting_new_value)
async def select_datetime_for_edit(callback: CallbackQuery, state: FSMContext, session: AsyncSession) -> None:
    """Handle date selection for event editing"""
    if callback.data is None:
        return
    
    date_str = callback.data.split(":")[1]  # Format: YYYYMMDDHHMM
    try:
        dt = datetime.strptime(date_str, "%Y%m%d%H%M")
    except ValueError:
        await callback.answer("Неверный формат даты", show_alert=True)
        return

    data = await state.get_data()
    event_id = data.get("event_id")
    
    if event_id is None:
        await callback.answer("Ошибка: событие не выбрано", show_alert=True)
        return

    await repo.update_training(session, event_id, dt=dt)
    await state.clear()
    await callback.message.answer(f"Обновлено ✅\nНовая дата: {dt.strftime(DATE_FORMAT_DISPLAY)}")
    await callback.answer()


@router.message(AdminEventStates.waiting_new_value)
async def input_new_value(message: Message, state: FSMContext, session: AsyncSession) -> None:
    if message.text is None:
        return
    
    data = await state.get_data()
    event_id = data.get("event_id")
    field = data.get("field")
    new_value: str | datetime = message.text

    if event_id is None or field is None:
        await message.answer("Ошибка: поле не выбрано")
        await state.clear()
        return

    training = await repo.get_training_by_id(session, event_id)
    if training is None:
        await message.answer("Событие не найдено")
        await state.clear()
        return

    # Map field names to database column names
    field_mapping: dict[str, str] = {
        "title": "title",
        "datetime": "dt",
        "place": "location",
        "description": "description",
    }

    db_field = field_mapping[field] if field in field_mapping else field

    # Convert datetime if needed
    if db_field == "dt":
        if isinstance(new_value, str):
            dt = _parse_datetime_manual(new_value)
            if dt is None:
                await message.answer("Неверный формат даты. Используйте формат ДД.ММ ЧЧ:ММ (например: 15.01 14:00)")
                return
            new_value = dt
        # If already datetime, use as-is

    await repo.update_training(session, event_id, **{db_field: new_value})
    await state.clear()
    await message.answer("Обновлено ✅")


# ── UC-08 Delete event ──────────────────────────────────────────────────────

@router.callback_query(F.data == "admin_delete_event")
async def cmd_delete_event(callback: CallbackQuery, session: AsyncSession) -> None:
    trainings = await repo.get_upcoming_trainings(session)

    if not trainings:
        await callback.message.answer("Нет предстоящих событий для удаления")
        await callback.answer()
        return

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t.title, callback_data=f"delete_select:{t.id}")]
        for t in trainings
    ])

    await callback.message.answer("Выберите событие для удаления:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("delete_select:"))
async def select_event_to_delete(callback: CallbackQuery, session: AsyncSession) -> None:
    if callback.data is None:
        return
    event_id = int(callback.data.split(":")[1])

    training = await repo.get_training_by_id(session, event_id)
    if training is None:
        await callback.message.answer("Событие не найдено")
        await callback.answer()
        return

    deleted = await repo.delete_training(session, event_id)
    if deleted:
        await callback.message.answer("Удалено ✅")
    else:
        await callback.message.answer("Не удалось удалить событие (возможно, оно уже прошло)")
    await callback.answer()


# ── UC-09 Export participants ───────────────────────────────────────────────

@router.callback_query(F.data == "admin_export_participants")
async def cmd_export_participants(callback: CallbackQuery, session: AsyncSession) -> None:
    trainings = await repo.get_upcoming_trainings(session)

    if not trainings:
        await callback.message.answer("Нет событий для экспорта участников")
        await callback.answer()
        return

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t.title, callback_data=f"export_select:{t.id}")]
        for t in trainings
    ])

    await callback.message.answer("Выберите событие для экспорта участников:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("export_select:"))
async def select_event_to_export(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.data is None:
        return
    event_id = int(callback.data.split(":")[1])
    await state.update_data(event_id=event_id)

    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="XLSX", callback_data="export_format:xlsx")],
        [InlineKeyboardButton(text="DOCX", callback_data="export_format:docx")],
    ])

    await callback.message.answer("Выберите формат экспорта:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("export_format:"))
async def select_export_format(callback: CallbackQuery, state: FSMContext, session: AsyncSession) -> None:
    if callback.data is None:
        return
    fmt = callback.data.split(":")[1]
    data = await state.get_data()
    event_id = data.get("event_id")
    
    if event_id is None:
        await callback.answer("Ошибка: событие не выбрано", show_alert=True)
        return

    participants = await repo.get_training_participants(session, event_id)

    if not participants:
        await callback.message.answer("Нет участников")
        await state.clear()
        await callback.answer()
        return

    # Generate file based on format
    if fmt == "xlsx":
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Participants"
        ws.append(["ID", "Full Name", "Username", "Email", "Phone"])
        for p in participants:
            ws.append([p.id, p.full_name, p.username or "", p.email or "", p.phone or ""])

        from io import BytesIO
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        await callback.message.answer_document(
            types.BufferedInputFile(
                file=buffer.getvalue(),
                filename=f"participants_{event_id}.xlsx"
            ),
            caption=f"Участники события {event_id}"
        )
    else:  # docx
        from docx import Document
        doc = Document()
        doc.add_heading(f"Участники события {event_id}", 0)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Full Name'
        hdr_cells[2].text = 'Username'
        hdr_cells[3].text = 'Email'
        hdr_cells[4].text = 'Phone'

        for p in participants:
            row_cells = table.add_row().cells
            row_cells[0].text = str(p.id)
            row_cells[1].text = p.full_name
            row_cells[2].text = p.username or ""
            row_cells[3].text = p.email or ""
            row_cells[4].text = p.phone or ""

        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        await callback.message.answer_document(
            types.BufferedInputFile(
                file=buffer.getvalue(),
                filename=f"participants_{event_id}.docx"
            ),
            caption=f"Участники события {event_id}"
        )

    await state.clear()
    await callback.answer()


# ── UC-10 Broadcast ─────────────────────────────────────────────────────────

@router.callback_query(F.data == "admin_broadcast")
async def cmd_broadcast(callback: CallbackQuery, state: FSMContext) -> None:
    await state.set_state(AdminBroadcastStates.waiting_text)
    await callback.message.answer("Введите текст рассылки:")
    await callback.answer()


@router.message(AdminBroadcastStates.waiting_text)
async def input_broadcast_text(message: Message, state: FSMContext, session: AsyncSession, bot: Bot) -> None:
    if message.text is None:
        return
    
    text = message.text

    users = await repo.get_all_users(session)

    if not users:
        await message.answer("Нет пользователей для рассылки")
        await state.clear()
        return

    sent_count = 0
    for user in users:
        try:
            await bot.send_message(user.tg_id, text)
            sent_count += 1
        except Exception as e:
            # Log error but continue
            print(f"Failed to send message to user {user.tg_id}: {e}")

    await state.clear()
    await message.answer(f"Отправлено {sent_count} пользователям ✅")


# ── UC-11 Manage admins ─────────────────────────────────────────────────────

@router.callback_query(F.data == "admin_manage_admins")
async def cmd_manage_admins(callback: CallbackQuery, state: FSMContext) -> None:
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Назначить админа", callback_data="admin_action:grant")],
        [InlineKeyboardButton(text="Снять админа", callback_data="admin_action:revoke")],
    ])

    await callback.message.answer("Выберите действие:", reply_markup=keyboard)
    await callback.answer()


@router.callback_query(F.data.startswith("admin_action:"))
async def select_admin_action(callback: CallbackQuery, state: FSMContext) -> None:
    if callback.data is None:
        return
    action = callback.data.split(":")[1]
    await state.update_data(action=action)
    await state.set_state(AdminManageAdminStates.waiting_user_id)
    await callback.message.answer("Введите Telegram ID пользователя:")
    await callback.answer()


@router.message(AdminManageAdminStates.waiting_user_id)
async def input_user_id(message: Message, state: FSMContext, session: AsyncSession) -> None:
    if message.text is None or message.from_user is None:
        return
    
    data = await state.get_data()
    action = data.get("action")

    try:
        user_id = int(message.text)
    except ValueError:
        await message.answer("Неверный формат ID. Введите число.")
        return

    # Check if trying to modify self
    if user_id == message.from_user.id:
        await message.answer("Нельзя изменить права самому себе")
        await state.clear()
        return

    # Find user by tg_id
    user = await repo.get_user_by_tg_id(session, user_id)
    if user is None:
        await message.answer("Пользователь не найден")
        await state.clear()
        return

    # Update admin status
    is_admin = (action == "grant")
    await repo.update_user(session, user_id, is_admin=is_admin)

    if is_admin:
        await message.answer("Назначен ✅")
    else:
        await message.answer("Права сняты ✅")

    await state.clear()
